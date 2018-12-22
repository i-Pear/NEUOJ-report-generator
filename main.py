#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from urllib import request
from urllib.request import urlopen
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches

name='王一蒙'
username='20184366'
username1=20184366

path="template/{0}.docx"
document = Document(path.format(1))  # 打开文件.docx
url='https://oj.neu.edu.cn/status/p/1?username={0}&pid={1}&lang=All&result=Accepted'
url1='https://oj.neu.edu.cn/status/{0}'

for pid in range(652, 705):
    htmlfile = urlopen(url.format(username,pid))
    htmlhandle = htmlfile.read()
    bsObj = BeautifulSoup(htmlhandle,'html.parser')
    id=''
    for herolist in bsObj.findAll("tr",{"class":"front-table-row"}):
        id=herolist.td.attrs['title']
        title=herolist.find("td",{"class":"text-left"}).find("a").get_text()
        print(title + "   https://oj.neu.edu.cn/status/" + id)
        p = document.add_paragraph()
        run = p.add_run('      C语言程序设计实验报告')
        p.bold = True
        font = run.font
        font.name = u'宋体'
        font.size = Pt(24)
        p = document.add_paragraph()
        run = p.add_run("班级： 计算机1805 姓名： {0}  学号： {1} ".format(name, username1))
        run.bold = True
        font = run.font
        font.name = u'宋体'
        font.size = Pt(16)
        p = document.add_paragraph()
        run = p.add_run("实验题目：{0}".format(title))
        run.bold = True
        font = run.font
        font.name = u'宋体'
        font.size = Pt(16)
        p = document.add_paragraph('     ')
        run = p.add_run(("对应源代码:"))
        run.bold = True
        font = run.font
        font.name = u'宋体'
        font.size = Pt(16)
        p = document.add_paragraph()
        run = p.add_run(("https://oj.neu.edu.cn/status/{0}".format(id)))
        run.bold = True
        font = run.font
        font.name = u'宋体'
        font.size = Pt(8)
        p = document.add_paragraph('   ')
        run = p.add_run("实验结果:")
        run.bold = True
        font = run.font
        font.name = u'宋体'
        font.size = Pt(16)
        p = document.add_paragraph('   ')
        document.add_page_break()
document.save(path.format(username1))
